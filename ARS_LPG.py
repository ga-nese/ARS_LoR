import streamlit as st
from fpdf import FPDF
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches  # <-- NEW IMPORT ADDED HERE
from io import BytesIO
import datetime
import os

# --- CONFIGURE THE PAGE ---
st.set_page_config(page_title="ARS LoR_LP", layout="centered")

st.title("ARS LetterPad Plugin")
st.markdown("""
Upload a **Word Document (.docx)** containing the body of the letter. 
The app will generate a formatted PDF on the official letterhead, as well as an editable Word version.
""")

# --- PDF CLASS DEFINITION ---
class PDF(FPDF):
    def header(self):
        if self.page_no() == 1:
            # 1. LOGO AND UNIVERSITY HEADER
            header_path = "header.png" 
            
            if os.path.exists(header_path):
                self.image(header_path, x=10, y=8, w=190)
                self.ln(34) 
            else:
                self.set_text_color(60, 120, 60) 
                self.set_font('Times', 'B', 16)
                self.cell(0, 5, 'KERALA AGRICULTURAL UNIVERSITY', 0, 1, 'C')
                self.ln(10)

            # 2. PROFESSOR DETAILS
            self.set_text_color(0, 0, 0) 
            self.set_font('Times', 'B', 11)
            
            right_margin_start = 110
            self.set_xy(right_margin_start, self.get_y())
            self.cell(0, 5, 'Dr. Archana Raghavan Sathyan', 0, 1, 'R')
            
            self.set_font('Times', 'I', 11)
            self.set_xy(right_margin_start, self.get_y())
            self.cell(0, 5, 'DAAD Research Ambassador', 0, 1, 'R')
            
            self.set_xy(right_margin_start, self.get_y())
            self.cell(0, 5, 'Assistant Professor', 0, 1, 'R')
            
            self.set_xy(right_margin_start, self.get_y())
            self.cell(0, 5, 'Department of Agricultural Extension Education', 0, 1, 'R')

            self.set_xy(right_margin_start, self.get_y())
            self.cell(0, 5, 'Email: archana.rs@kau.in', 0, 1, 'R')
            
            self.ln(2) 
        else:
            self.ln(20)

    def footer(self):
        self.set_y(-15)
        self.set_font('Arial', 'I', 8)
        self.set_text_color(0, 0, 0)
        self.cell(0, 10, 'Page ' + str(self.page_no()) + '/{nb}', 0, 0, 'C')

# --- MAIN APP LOGIC ---

is_lor = st.radio("Is this a recommendation letter?", options=["Yes", "No"])
uploaded_file = st.file_uploader("Choose your Word file", type="docx")

if uploaded_file is not None:
    # 1. Read the Word Document
    doc = Document(uploaded_file)
    full_text = []
    
    # --- TEXT CLEANING FUNCTION ---
    def clean_text(text):
        replacements = {
            "\u2018": "'", "\u2019": "'", "\u201c": '"', "\u201d": '"', 
            "\u2013": "-", "\u2014": "-"
        }
        for original, replacement in replacements.items():
            text = text.replace(original, replacement)
        return text.encode('latin-1', 'replace').decode('latin-1')

    for para in doc.paragraphs:
        clean_para = clean_text(para.text)
        full_text.append(clean_para)
    
    # ---------------------------------------------------------
    # 2. GENERATE PDF
    # ---------------------------------------------------------
    pdf = PDF()
    pdf.alias_nb_pages()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    
    if is_lor == "Yes":
        pdf.set_font('Times', 'B', 12)
        pdf.set_text_color(0, 0, 0) 
        pdf.cell(0, 10, 'LETTER OF RECOMMENDATION', 0, 1, 'C')
        pdf.ln(2) 
    else:
        pdf.ln(12) 
    
    pdf.set_font('Times', '', 11)
    
    for para in full_text:
        if para.strip(): 
            pdf.multi_cell(0, 5, para)
            pdf.ln(2) 
            
    pdf.ln(5) 
    now = datetime.datetime.now()
    date_str = now.strftime("%d-%m-%Y") 
    
    if pdf.get_y() > 250:
        pdf.add_page()
        
    y_pos = pdf.get_y()
    
    pdf.set_xy(10, y_pos)
    pdf.cell(50, 5, 'Vellayani', 0, 1, 'L')
    pdf.set_xy(10, y_pos + 6) 
    pdf.cell(50, 5, date_str, 0, 1, 'L')
    
    pdf.set_xy(10, y_pos) 
    pdf.cell(0, 5, 'Yours sincerely,', 0, 1, 'R')
    pdf.set_xy(10, y_pos + 20) 
    pdf.cell(0, 5, 'Dr. Archana Raghavan Sathyan', 0, 1, 'R')

    pdf_output = pdf.output(dest='S').encode('latin-1', 'ignore') 
    
    # ---------------------------------------------------------
    # 3. GENERATE WORD (.DOCX)
    # ---------------------------------------------------------
    out_doc = Document()
    
    # --- NEW: Image Header Logic for Word ---
    header_path = "header.png"
    if os.path.exists(header_path):
        p_img = out_doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_img = p_img.add_run()
        # Insert the image and scale it to fit standard Word margins
        r_img.add_picture(header_path, width=Inches(6.0)) 
    else:
        # Fallback Text Header if image is missing
        h1 = out_doc.add_paragraph("KERALA AGRICULTURAL UNIVERSITY")
        h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        h1.runs[0].bold = True

    # Professor Details
    prof_details = [
        "Dr. Archana Raghavan Sathyan",
        "DAAD Research Ambassador",
        "Assistant Professor",
        "Department of Agricultural Extension Education",
        "Email: archana.rs@kau.in"
    ]
    for detail in prof_details:
        p = out_doc.add_paragraph(detail)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if detail == "Dr. Archana Raghavan Sathyan":
            p.runs[0].bold = True

    out_doc.add_paragraph() # Spacer

    # Conditional Title
    if is_lor == "Yes":
        p_title = out_doc.add_paragraph("LETTER OF RECOMMENDATION")
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_title.runs[0].bold = True

    out_doc.add_paragraph() # Spacer

    # Body Paragraphs
    for para in full_text:
        if para.strip():
            out_doc.add_paragraph(para)

    # Footer / Sign-off
    out_doc.add_paragraph() # Spacer
    p_place = out_doc.add_paragraph(f"Vellayani\n{date_str}")
    
    p_sign = out_doc.add_paragraph("Yours sincerely,\n\n\nDr. Archana Raghavan Sathyan")
    p_sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Save to memory buffer
    word_buffer = BytesIO()
    out_doc.save(word_buffer)
    word_buffer.seek(0)

    # ---------------------------------------------------------
    # 4. RENDER DOWNLOAD BUTTONS
    # ---------------------------------------------------------
    st.success("Files Generated Successfully!")
    
    file_prefix = "Recommendation" if is_lor == "Yes" else "Letter"
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.download_button(
            label="📄 Download PDF",
            data=pdf_output,
            file_name=f"{file_prefix}_{date_str}.pdf",
            mime="application/pdf"
        )
        
    with col2:
        st.download_button(
            label="📝 Download Word Document",
            data=word_buffer,
            file_name=f"{file_prefix}_{date_str}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
